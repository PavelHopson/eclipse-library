"""Build the illustrated guide and a readable DOCX from reviewed local sources."""
from pathlib import Path
import re
from html import escape
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / 'web/experiments/gwen-reveal'
GUIDE = ROOT / 'guides/gwen-stacy-canvas-reveal.md'

def inline(s):
    s = escape(s)
    s = re.sub(r'\[([^\]]+)\]\((https://[^\s)]+)\)', r'<a href="\2" target="_blank" rel="noopener noreferrer">\1</a>', s)
    s = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', s)
    return re.sub(r'`([^`]+)`', r'<code>\1</code>', s)

def blocks(text):
    lines = text.splitlines(); i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip(): i += 1; continue
        if line.startswith('```'):
            code = []; i += 1
            while i < len(lines) and not lines[i].startswith('```'): code.append(lines[i]); i += 1
            i += 1; yield ('code', '\n'.join(code)); continue
        if line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                if not re.fullmatch(r'[| :\-]+', lines[i]): rows.append([c.strip() for c in lines[i].strip('|').split('|')])
                i += 1
            yield ('table', rows); continue
        m = re.match(r'^(#{1,3}) (.+)', line)
        if m: yield ('h'+str(len(m[1])), m[2]); i += 1; continue
        if line.startswith('- '): yield ('li', line[2:]); i += 1; continue
        if line.startswith('> '): yield ('quote', line[2:]); i += 1; continue
        yield ('p', line); i += 1

def shade(cell, color):
    el = OxmlElement('w:shd'); el.set(qn('w:fill'), color); cell._tc.get_or_add_tcPr().append(el)

def runs(p, text):
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    for token in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        r = p.add_run(token.strip('*`') if token.startswith(('**', '`')) else token)
        if token.startswith('**'): r.bold = True
        if token.startswith('`'): r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string('982952')

def build():
    content = GUIDE.read_text(encoding='utf-8')
    original = (PROJECT/'assets/original-brief.txt').read_text(encoding='utf-8')
    parsed = list(blocks(content))
    headings = [(i, t) for i,(kind,t) in enumerate(parsed) if kind == 'h2']
    toc = ''.join(f'<a href="#section-{i}">{inline(t)}</a>' for i,t in headings)
    body = []
    for i,(kind,text) in enumerate(parsed):
        if kind == 'h1': continue
        if kind in ('h2','h3'): body.append(f'<{kind} id="section-{i}">{inline(text)}</{kind}>')
        elif kind == 'code': body.append('<pre><code>'+escape(text)+'</code></pre>')
        elif kind == 'table': body.append('<div class="table-scroll"><table>'+''.join('<tr>'+''.join(f'<{"th" if j==0 else "td"}>{inline(c)}</{"th" if j==0 else "td"}>' for c in row)+'</tr>' for j,row in enumerate(text))+'</table></div>')
        elif kind == 'li': body.append('<ul><li>'+inline(text)+'</li></ul>')
        elif kind == 'quote': body.append('<blockquote>'+inline(text)+'</blockquote>')
        else: body.append('<p>'+inline(text)+'</p>')
    html = '''<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><meta name="robots" content="noindex,nofollow"><meta http-equiv="Content-Security-Policy" content="default-src 'none'; style-src 'self'; img-src 'self'; font-src 'self'; base-uri 'none'; form-action 'none'"><title>Gwen: по ту сторону маски • Гайд</title><link rel="stylesheet" href="guide.css"></head><body>
<header class="guide-bar"><a href="../../index.html#guidesFeat">Eclipse Library</a><a href="index.html">Открыть эксперимент ↗</a></header>
<section class="cover"><div class="cover-text"><p class="edition">CANVAS / ИНТЕРАКТИВНЫЙ ДИЗАЙН</p><h1>По ту сторону<br><span>маски.</span></h1><p>Как два портрета превращаются<br>в живой комикс-постер.</p><div class="cover-actions"><a class="primary" href="index.html">Попробовать эффект ↗</a><a href="assets/Gwen-Stacy-Guide.docx" download>Скачать гайд DOCX ↓</a></div></div><figure><img src="assets/source-pair.webp" alt="Исходные изображения: Гвен в костюме и портрет без маски"><figcaption>Два предоставленных кадра. Общая геометрия. Никакой генерации лица.</figcaption></figure></section>
<div class="reading-layout"><nav class="contents" aria-label="Содержание"><strong>В этом гайде</strong>'''+toc+'''<a href="#full-brief">Полный исходный промпт</a></nav><main class="article">'''+''.join(body)+'''
<h2 id="full-brief">Полный исходный промпт</h2><p>Текст из предоставленного документа. Технический разбор и дополнения реализации находятся выше.</p><details class="full-brief"><summary>Открыть полный текст промпта</summary><pre>'''+escape(original)+'''</pre></details>
<footer class="article-end"><p>Материалы для локального прототипа. Публичное размещение и права на изображения проверяются отдельно.</p><a href="index.html">Вернуться к эффекту ↗</a></footer></main></div></body></html>'''
    (PROJECT/'guide.html').write_text(html, encoding='utf-8')

    doc = Document()
    sec = doc.sections[0]; sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
    sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.65); sec.left_margin = sec.right_margin = Inches(.78)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string('28252F'); normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [('Title',42,'18141D'),('Heading 1',23,'8D1E50'),('Heading 2',16,'28252F'),('Heading 3',12,'28252F')]:
        style = doc.styles[name]; style.font.name = 'Calibri'; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(18); style.paragraph_format.space_after = Pt(10)
    header = sec.header.paragraphs[0]; header.text = 'ECLIPSE LIBRARY  /  CANVAS STUDIO'
    header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor.from_string('8D1E50')
    foot = sec.footer.paragraphs[0]; foot.text = 'GWEN  •  Интерактивный фан-концепт                                       '
    foot.runs[0].font.size = Pt(8)
    field = OxmlElement('w:fldSimple'); field.set(qn('w:instr'), 'PAGE'); foot._p.append(field)
    doc.add_paragraph('ПРАКТИЧЕСКИЙ ГАЙД', 'Subtitle')
    doc.add_paragraph('По ту сторону\nмаски.', 'Title')
    p = doc.add_paragraph('Gwen Stacy / Earth-65\nДва портрета. Одна капля. Живое раскрытие.'); p.runs[0].font.size = Pt(17)
    # python-docx accepts PNG/JPEG, not WebP. A local generated PNG preserves the same pixels.
    from PIL import Image
    proof_png = ROOT/'.artifacts/gwen/source-pair.png'; proof_png.parent.mkdir(parents=True, exist_ok=True)
    Image.open(PROJECT/'assets/source-pair.webp').save(proof_png)
    doc.add_picture(str(proof_png), width=Inches(6.65))
    doc.add_paragraph('HTML + CSS + vanilla JavaScript\nCanvas 2D • 6 модулей • 12 уроков • полный исходный бриф', 'Subtitle')
    doc.add_paragraph('Подготовлено для Eclipse Library. 5 сентября 2026 года.\nИсходники предоставлены пользователем. Неофициальный фан-концепт; права на публичное распространение изображений отдельно не подтверждены.')
    doc.add_page_break()
    doc.add_heading('Содержание', level=1)
    for _, title in headings: doc.add_paragraph(title)
    doc.add_paragraph('Полный исходный промпт включён в приложение. Текст из исходного документа отделён от инженерных дополнений.')
    doc.add_page_break()
    for kind,text in parsed:
        if kind == 'h1': continue
        if kind == 'h2':
            p=doc.add_heading(text, level=1); p.paragraph_format.page_break_before=True
        elif kind == 'h3': doc.add_heading(text, level=2)
        elif kind == 'table':
            table=doc.add_table(rows=0, cols=len(text[0])); table.autofit=True
            for i,row in enumerate(text):
                cells=table.add_row().cells
                for cell,value in zip(cells,row):
                    runs(cell.paragraphs[0], value)
                    shade(cell,'EDE8F0' if i==0 else ('F8F6F9' if i%2 else 'FFFFFF'))
                    for r in cell.paragraphs[0].runs: r.font.size=Pt(9); r.bold=(i==0)
                if i==0:
                    repeat=OxmlElement('w:tblHeader'); table.rows[i]._tr.get_or_add_trPr().append(repeat)
            doc.add_paragraph()
        elif kind == 'code':
            for line in text.splitlines():
                p=doc.add_paragraph(); r=p.add_run(line or ' '); r.font.name='Consolas'; r.font.size=Pt(8)
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2EEF4'); p._p.get_or_add_pPr().append(shd)
            doc.add_paragraph()
        else:
            p=doc.add_paragraph(style='List Bullet' if kind=='li' else None); runs(p,text)
            if kind=='quote':
                p.paragraph_format.left_indent=Inches(.18)
                for r in p.runs: r.italic=True; r.font.color.rgb=RGBColor.from_string('8D1E50')
    p=doc.add_heading('Полный исходный промпт', level=1); p.paragraph_format.page_break_before=True
    doc.add_paragraph('Ниже сохранён полный текст предоставленного документа. Изображения показаны на обложке. Оформление изменено, технические формулировки и значения не сокращены.')
    for line in original.splitlines():
        if not line.strip(): continue
        if line.isupper() and len(line)<100 and not re.search(r'[#<>;={}]', line): doc.add_heading(line, level=2)
        else:
            p=doc.add_paragraph(line); p.paragraph_format.space_after=Pt(2)
            p.paragraph_format.line_spacing=1.08
            for r in p.runs: r.font.size=Pt(9)
    doc.core_properties.title='Gwen: по ту сторону маски'
    doc.core_properties.subject='Canvas reveal: практический гайд и полный исходный бриф'
    doc.core_properties.author='Eclipse Library'
    path=PROJECT/'assets/Gwen-Stacy-Guide.docx'; doc.save(path)
    print(f'Built guide.html and {path.name}; {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables. Original brief: {len(original)} characters.')

if __name__ == '__main__': build()

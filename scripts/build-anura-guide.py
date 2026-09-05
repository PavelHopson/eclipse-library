"""Build a clean, illustrated HTML/DOCX guide without the source DOCX's OLE package."""
from pathlib import Path
import importlib.util, sys
sys.dont_write_bytecode = True
from html import escape
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
DEST = ROOT / 'web/experiments/anura'
spec = importlib.util.spec_from_file_location('guide_common', ROOT/'scripts/build-gwen-guide.py')
common = importlib.util.module_from_spec(spec); spec.loader.exec_module(common)

def build():
    original=(ROOT/'.artifacts/anura/original-brief.txt').read_text(encoding='utf8')
    (DEST/'assets/original-brief.txt').write_text(original,encoding='utf8')
    parsed=list(common.blocks((ROOT/'guides/anura-interactive-frog.md').read_text(encoding='utf8')))
    headings=[(i,t) for i,(kind,t) in enumerate(parsed) if kind=='h2']
    toc=''.join(f'<a href="#section-{i}">{common.inline(t)}</a>' for i,t in headings)
    body=[]
    for i,(kind,text) in enumerate(parsed):
        if kind=='h1': continue
        if kind in ('h2','h3'): body.append(f'<{kind} id="section-{i}">{common.inline(text)}</{kind}>')
        elif kind=='code': body.append('<pre><code>'+escape(text)+'</code></pre>')
        elif kind=='li': body.append('<ul><li>'+common.inline(text)+'</li></ul>')
        elif kind=='quote': body.append('<blockquote>'+common.inline(text)+'</blockquote>')
        else: body.append('<p>'+common.inline(text)+'</p>')
    html='''<!doctype html><html lang="ru"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><meta name="robots" content="noindex,nofollow"><meta http-equiv="Content-Security-Policy" content="default-src 'none'; style-src 'self'; img-src 'self'; font-src 'self'; base-uri 'none'; form-action 'none'"><title>ANURA — взгляд в ответ • Гайд</title><link rel="stylesheet" href="guide.css"></head><body>
<header class="guide-bar"><a href="../../animations.html#anura-experiment">← Eclipse Library</a><a href="index.html">Открыть ANURA ↗</a></header>
<section class="cover"><div class="cover-text"><p class="edition">FIELD NOTES / 01 · CANVAS EXPERIENCE</p><h1>Взгляд<br><span>в ответ.</span></h1><p>Как короткое видео становится<br>интерактивным персонажем.</p><div class="cover-actions"><a class="primary" href="index.html">Поймать взгляд ↗</a><a href="assets/ANURA-Guide.docx" download>Скачать гайд DOCX ↓</a></div></div><figure><img src="assets/source.webp" alt="Исходное изображение: зелёная лягушка с большими глазами на бирюзовом студийном фоне"><figcaption>Исходник пользователя. Без генеративной замены персонажа.</figcaption></figure></section>
<section class="sequence-strip" aria-label="Направления взгляда"><img src="assets/look-directions.webp" alt="Четыре настоящих кадра видео: взгляд влево, вверх, вправо и на зрителя"><p>10 секунд исходного видео → 240 WebP → один управляемый взгляд.</p></section>
<div class="reading-layout"><nav class="contents" aria-label="Содержание"><strong>Полевые заметки</strong>'''+toc+'''<a href="#full-brief">Полный исходный промпт</a></nav><main class="article">'''+''.join(body)+'''
<h2 id="full-brief">Полный исходный промпт</h2><p>Все 98 непустых абзацев из документа пользователя. Исходные требования отделены от фактических решений выше; вложенный OLE-объект исключён.</p><details class="full-brief"><summary>Открыть полный текст: изображение, видео и сайт</summary><pre>'''+escape(original)+'''</pre></details><footer class="article-end"><p>Подготовлено для Eclipse Library · 5 сентября 2026. Права на публичное распространение исходников проверяются отдельно.</p><a href="index.html">Вернуться к лягушке ↗</a></footer></main></div></body></html>'''
    (DEST/'guide.html').write_text(html,encoding='utf8')
    strip=Image.new('RGB',(1600,260),'#17352f')
    for i,index in enumerate([30,84,138,0]):
        with Image.open(DEST/f'frames/{index:03}.webp') as frame:
            # Same crop for every pose; no synthesis or alteration of the character.
            tile=frame.crop((620,70,1250,700)); tile.thumbnail((260,260))
            strip.paste(tile,(i*400+70,0))
    strip.save(DEST/'assets/look-directions.webp',quality=83)
    artwork=ROOT/'.artifacts/anura'; strip.save(artwork/'look-directions.png')
    with Image.open(DEST/'assets/source.webp') as source: source.save(artwork/'guide-cover.png')
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.27);sec.page_height=Inches(11.69)
    sec.top_margin=sec.bottom_margin=Inches(.7);sec.left_margin=sec.right_margin=Inches(.8)
    normal=doc.styles['Normal'];normal.font.name='Calibri';normal.font.size=Pt(10.5)
    normal.font.color.rgb=RGBColor.from_string('233C35');normal.paragraph_format.space_after=Pt(8)
    normal.paragraph_format.line_spacing=1.12
    for name,size,color in [('Title',48,'153E33'),('Heading 1',25,'285C47'),('Heading 2',16,'243C32')]:
        s=doc.styles[name];s.font.name='Calibri';s.font.size=Pt(size);s.font.color.rgb=RGBColor.from_string(color)
        s.paragraph_format.space_before=Pt(18);s.paragraph_format.space_after=Pt(10)
    sec.header.paragraphs[0].text='ANURA®   /   FIELD NOTES 01                                               ECLIPSE LIBRARY'
    for r in sec.header.paragraphs[0].runs:r.font.size=Pt(8)
    foot=sec.footer.paragraphs[0];foot.text='ПРАКТИЧЕСКИЙ ГАЙД  •  2026                                                   '
    f=OxmlElement('w:fldSimple');f.set(qn('w:instr'),'PAGE');foot._p.append(f)
    doc.add_paragraph('CANVAS / ИНТЕРАКТИВНЫЙ ДИЗАЙН','Subtitle')
    doc.add_paragraph('Взгляд\nв ответ.','Title')
    doc.add_paragraph('ANURA — Interactive Frog\nОт короткого ролика к живому наблюдению.','Subtitle')
    doc.add_picture(str(artwork/'guide-cover.png'),width=Inches(6.65))
    doc.add_paragraph('240 кадров  /  4,5 МБ  /  HTML + CSS + JavaScript')
    doc.add_paragraph('6 модулей • 12 уроков • полный исходный промпт\nИсходники предоставлены пользователем. 5 сентября 2026. Права на распространение отдельно не подтверждены.')
    doc.add_page_break();doc.add_heading('Карта наблюдения',1)
    for _,title in headings:doc.add_paragraph(title)
    doc.add_picture(str(artwork/'look-directions.png'),width=Inches(6.65))
    doc.add_paragraph('Четыре кадра исходного видео. Реакция выбирает существующее изображение, а не генерирует новое.')
    for kind,text in parsed:
        if kind=='h1':continue
        if kind=='h2':
            p=doc.add_heading(text,1);p.paragraph_format.page_break_before=True
        elif kind=='h3':doc.add_heading(text,2)
        elif kind=='code':
            for line in text.splitlines():
                p=doc.add_paragraph(line);p.paragraph_format.space_after=Pt(2)
                for r in p.runs:r.font.name='Consolas';r.font.size=Pt(9)
        else:
            p=doc.add_paragraph(style='List Bullet' if kind=='li' else None);common.runs(p,text)
            if kind=='quote':
                for r in p.runs:r.italic=True
    p=doc.add_heading('Приложение. Полный исходный промпт',1);p.paragraph_format.page_break_before=True
    doc.add_paragraph('Текст сохранён целиком. OLE-вложение не переносилось; видео и исходное изображение обработаны отдельно. Требования промпта — исходная задумка, технические уточнения находятся в модулях выше.')
    for line in original.splitlines():
        if not line.strip():continue
        p=doc.add_paragraph(line);p.paragraph_format.space_after=Pt(4)
        for r in p.runs:r.font.size=Pt(9)
    doc.core_properties.title='ANURA — взгляд в ответ';doc.core_properties.author='Eclipse Library'
    doc.core_properties.subject='Интерактивный Canvas, видеокадры, полный исходный промпт'
    doc.save(DEST/'assets/ANURA-Guide.docx')
    print(f'ANURA guide: {len(doc.paragraphs)} paragraphs; {len(headings)} modules; original text {len(original)} characters.')

if __name__=='__main__':build()
